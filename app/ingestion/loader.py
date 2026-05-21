"""
Document loader for PDF, DOCX, and TXT files.

Extracts text content and attaches per-page metadata so that downstream
components can cite exact sources.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List

import docx
import fitz  # PyMuPDF

from langchain_core.documents import Document

from app.utils.helpers import get_file_extension

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load documents from supported file formats into LangChain Documents."""

    _SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def load(self, file_path: str) -> List[Document]:
        """
        Load a file and return a list of LangChain Document objects.

        Each Document carries metadata with at least:
        - ``source``: the original filename
        - ``page``: the 1-based page number (or 1 for single-page formats)

        Args:
            file_path: Absolute or relative path to the file.

        Returns:
            A list of :class:`Document` objects.

        Raises:
            FileNotFoundError: If *file_path* does not exist.
            ValueError: If the file extension is not supported.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = get_file_extension(path.name).lower()

        if ext not in self._SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format '{ext}'. "
                f"Supported formats: {', '.join(sorted(self._SUPPORTED_EXTENSIONS))}"
            )

        loader_map = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".txt": self._load_txt,
        }

        logger.info("Loading file: %s (format: %s)", path.name, ext)
        documents = loader_map[ext](path)
        logger.info("Loaded %d document segment(s) from %s", len(documents), path.name)
        return documents

    @staticmethod
    def _load_pdf(path: Path) -> List[Document]:
        """Extract text from a PDF file, one Document per page."""
        documents: List[Document] = []
        try:
            with fitz.open(str(path)) as pdf:
                for page_num, page in enumerate(pdf, start=1):
                    text = page.get_text("text")
                    if text.strip():
                        documents.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "source": path.name,
                                    "page": page_num,
                                },
                            )
                        )
        except Exception as exc:
            raise RuntimeError(
                f"Failed to parse PDF '{path.name}': {exc}"
            ) from exc
        return documents

    @staticmethod
    def _load_docx(path: Path) -> List[Document]:
        """Extract text from a DOCX file as a single Document."""
        try:
            doc = docx.Document(str(path))
            full_text = "\n".join(
                paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
            )
        except Exception as exc:
            raise RuntimeError(
                f"Failed to parse DOCX '{path.name}': {exc}"
            ) from exc

        if not full_text.strip():
            return []

        return [
            Document(
                page_content=full_text,
                metadata={"source": path.name, "page": 1},
            )
        ]

    @staticmethod
    def _load_txt(path: Path) -> List[Document]:
        """Read a plain-text file as a single Document."""
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
        except Exception as exc:
            raise RuntimeError(
                f"Failed to read TXT '{path.name}': {exc}"
            ) from exc

        if not text.strip():
            return []

        return [
            Document(
                page_content=text,
                metadata={"source": path.name, "page": 1},
            )
        ]
